from pathlib import Path
import os

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font(run, size=10.5, name="宋体", bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def write_para(paragraph, text, size=11, bold=False, align=None):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if align is not None:
        paragraph.alignment = align


def write_cell(cell, text, size=9.5, font="宋体", bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    for index, line in enumerate(str(text).splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=size, name=font, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def write_code(cell, text):
    write_cell(cell, text.strip("\n"), size=8, font="Consolas")


def fill_table(table, date, place, task, problems, code, summary):
    write_cell(table.cell(0, 1), date, size=10.5)
    write_cell(table.cell(0, 3), place, size=10.5)
    write_cell(table.cell(1, 1), task, size=10.5)
    for idx, row_index in enumerate([4, 5, 6, 7, 8]):
        if idx < len(problems):
            write_cell(table.cell(row_index, 0), problems[idx][0])
            write_cell(table.cell(row_index, 2), problems[idx][1])
        else:
            write_cell(table.cell(row_index, 0), "")
            write_cell(table.cell(row_index, 2), "")
    write_code(table.cell(10, 0), code)
    write_cell(table.cell(len(table.rows) - 1, 0), summary, size=9.5)


def main():
    desktop = Path(os.environ.get("USERPROFILE", r"C:\Users\30587")) / "Desktop"
    source = next(path for path in desktop.glob("202431212*.docx") if " - " not in path.name)
    output = desktop / "202431212  王振宇 实训报告 - 项目贴合版.docx"
    doc = Document(str(source))

    for style_name in ["Normal"]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)

    replacements = {
        "项目名称：校园生活百事通助手": "项目名称：安徽交通职业技术学院智慧官网与 AI 招生服务系统",
        "实训日期：2026年6月17日 - 2026年6月21日": "实训日期：2026年6月17日 - 2026年6月21日",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            write_para(paragraph, replacements[text], size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif text == "1. 项目概述（200字）":
            write_para(paragraph, text, size=14, bold=True)
        elif text.startswith("本项目名称为"):
            write_para(
                paragraph,
                "本项目名称为“安徽交通职业技术学院智慧官网与 AI 招生服务系统”。系统以学校官网改版为基础，建设了前后端分离的校园门户、后台管理、招生管理、官网实时采集、学生与教务管理、AI 招生咨询和校园生活百事通等功能。前端采用 Vue 3 + Element Plus，后端采用 FastAPI + SQLAlchemy + MySQL，并接入 Chroma 向量库和智谱 GLM 大模型。项目解决了传统官网信息展示单一、后台维护分散、招生咨询依赖人工、校园生活规则查询不便等问题。系统支持首页轮播、二级页面、学院内页、英文官网、官网采集审核、学生课表成绩、招生 AI 问答、校园生活 RAG 问答、绩点计算、教学周查询以及数字人语音播报，形成了一个集展示、管理、问答和服务于一体的综合 AI 应用。",
            )
        elif text == "2. 技术架构图":
            write_para(paragraph, text, size=14, bold=True)
        elif text.startswith("技术架构图："):
            write_para(
                paragraph,
                "技术架构图：用户/管理员 -> Vue 3 前端门户与后台 -> FastAPI 接口层 -> 业务模块（官网内容、招生管理、学生管理、教务管理、官网采集、AI 助手） -> MySQL 数据库 / Chroma 向量库 / uploads 文件资源 -> 智谱 GLM、DeepSeek、星火等大模型服务 -> 返回页面展示、语音播报和后台审核结果。",
            )
        elif text.startswith("流程说明："):
            write_para(
                paragraph,
                "流程说明：访客进入官网后可浏览首页、二级栏目、学院内页和英文官网；管理员登录后台后可维护招生、官网、学生、教务和采集数据；AI 助手接收用户问题后，根据 provider 选择智谱 GLM 等模型，根据 assistant_mode 分流到招生咨询或校园生活百事通。招生问题会结合专业库、分数线、官网采集内容和本地知识库生成回答；校园生活问题会使用 campus_data.csv 和 Chroma 向量库进行 RAG 检索；回答完成后，前端数字人调用浏览器语音合成接口朗读结果。",
            )
        elif text == "4. 核心代码展示":
            write_para(paragraph, text, size=14, bold=True)
        elif text.startswith("核心代码"):
            write_para(
                paragraph,
                "核心代码一：AI 服务统一入口 ask_ai(db, payload)\n该函数根据 assistant_mode 判断用户使用的是招生咨询还是校园生活助手；根据 provider 选择智谱 GLM、星火、DeepSeek 或 OpenAI 兼容接口；先生成本地检索结果，再调用大模型进行回答。\n\n核心代码二：校园生活 RAG 检索 retrieve_campus_knowledge(question)\n该函数读取 campus_data.csv，结合 Chroma 向量库、同义词扩展和相似度重排，解决“发烧了怎么办”“灯坏了找谁”等口语问题无法直接命中制度关键词的问题。\n\n核心代码三：官网实时采集 crawler_service\n该模块负责从学校官网栏目抓取新闻、通知、招标采购、部门动态等内容，写入官网文章库，并在后台提供采集源、采集记录和审核入口。\n\n核心代码四：数字人朗读 speakAnswer(text)\n前端在 AI 回答生成完成后，调用浏览器 SpeechSynthesisUtterance 进行中文语音播报，同时通过 CSS 动画让数字人头像产生说话状态。",
            )
        elif text == "5. 运行截图":
            write_para(paragraph, text, size=14, bold=True)
        elif text.startswith("截图说明："):
            write_para(
                paragraph,
                "截图说明：\n1. 官网首页：展示安徽交通职业技术学院首页、轮播图、导航栏、新闻通知和专业集群。\n2. 后台管理：展示 ACVTC 管理后台、数据看板、招生管理、官网管理、官网采集等模块。\n3. 官网采集：展示采集源状态、采集记录、手动采集和进入官网审核功能。\n4. 学生/教务管理：展示学生名单、学生详情、课表、成绩和考勤管理。\n5. 招生 AI 助手：展示模型平台选择为“智谱 GLM”，输入问题后返回招生咨询回答。\n6. 数字人播报：展示 AI 回答完成后数字人头像处于朗读状态，并提供朗读、重播、停止按钮。\n7. 校园生活百事通：展示请假、奖学金、报修、一卡通、选课等 RAG 问答效果。\n提交 PDF 前可将以上功能页面截图插入本节。",
            )
        elif text == "6. 问题与反思（300字）":
            write_para(paragraph, text, size=14, bold=True)
        elif text.startswith("本次实训中遇到的最大困难"):
            write_para(
                paragraph,
                "本次实训中最大的困难是项目功能范围较大，不再是单一聊天机器人，而是包含官网门户、后台管理、学生教务、官网采集、AI 问答和数字人播报的完整系统。前端页面多、后端接口多、数据库模型多，任何一个模块的改动都可能影响整体运行。第二个困难是大模型平台接入。星火接口因为应用授权问题返回 11200，不能稳定调用，因此我改为接入智谱 GLM，并在系统中保留 Spark、DeepSeek、OpenAI 兼容平台选择，提升可替换性。第三个困难是中文编码和真实数据维护，Windows 环境下文件写入方式不正确时容易出现乱码。通过统一 UTF-8 写入、前端打包检查和关键词扫描，保证页面中文正常显示。如果后续继续完善，我希望增加更正式的权限体系、后台知识库维护界面、真实学校办事接口、语音合成服务端接口和更精细的日志监控，让系统更接近真实可部署产品。",
            )
        elif text == "7. 实训总结（200字）":
            write_para(paragraph, text, size=14, bold=True)
        elif text.startswith("通过本次实训"):
            write_para(
                paragraph,
                "通过本次实训，我对 AI 应用开发的理解从“单一模型问答”提升到了“完整系统工程”。一个可用的 AI 校园服务系统不仅需要大模型，还需要前端交互、后端接口、数据库、知识库、权限、采集、审核、日志和部署配置共同支撑。本项目中，我完成了 Vue 前端页面、FastAPI 后端接口、MySQL 数据管理、Chroma 检索增强、智谱 GLM 接入、校园生活 RAG、学生教务管理和数字人语音播报等功能。实践过程让我认识到，AI 应用的关键不是简单调用 API，而是把模型能力放进真实业务流程中，并通过本地数据、工具函数和兜底机制保证回答可信。此次项目提升了我的全栈开发能力、问题排查能力和 AI 工程化思维。",
            )

    day1_problems = [
        ("项目一开始只有官网展示需求，但后续扩展到后台、招生、学生、教务、AI 和采集，模块边界容易混乱。", "先梳理系统模块，将项目拆分为官网门户、后台管理、招生服务、学生教务、官网采集、AI 助手六条主线。"),
        ("旧版 Streamlit 单体应用不适合继续扩展官网和后台功能。", "删除旧单体思路，采用 Vue 3 + FastAPI 前后端分离结构，便于后续维护和部署。"),
        ("首页和二级页面如果只做静态页面，会和后台数据脱节。", "设计官网文章、栏目、轮播图、学院内页和服务页接口，让后台内容能够更新到前台。"),
        ("页面曾出现中文乱码和问号，影响展示效果。", "统一使用 UTF-8 写入源码和配置，打包前扫描明显乱码。"),
        ("项目内容较多，用户入口容易混乱。", "在首页、后台和服务页分别设置清晰入口，例如招生 AI 助手、学生课表、后台登录、官网采集。"),
    ]
    day1_code = r'''
// frontend/src/router/index.js
routes: [
  { path: '/', component: PortalView },
  { path: '/chat', component: ChatView },
  { path: '/schedule', component: StudentScheduleView },
  { path: '/channel/:name', component: ChannelView },
  { path: '/college/:name', component: CollegeDetailView },
  {
    path: '/admin',
    component: AdminLayout,
    meta: { requiresAuth: true },
    children: [
      { path: 'dashboard', component: DashboardView },
      { path: 'admissions', component: AdmissionsView },
      { path: 'website', component: WebsiteAdminView },
      { path: 'crawler', component: CrawlerAdminView },
      { path: 'students', component: StudentsView },
      { path: 'academic', component: AcademicView }
    ]
  }
]
'''

    day2_problems = [
        ("学生数据从 Excel 导入后，需要和后台学生管理、学生详情、课表成绩联动。", "扩展学生字段，保留学号、班级、联系方式、考试号等信息，并在学生详情页展示教务关联数据。"),
        ("学生登录和管理员登录容易混用。", "拆分管理员 token 和学生 token，学生端使用学号登录，只能查看自己的课表、成绩和考勤。"),
        ("课表需要按 18 周展示，并且课程分布固定。", "在教务模块中按周次、星期、节次组织课表数据，前端允许学生自由切换周次查看。"),
        ("成绩总评不能只手动填写。", "在教务管理中加入平时成绩、期末成绩、考核成绩和总评计算逻辑。"),
        ("后台权限如果只靠前端隐藏不安全。", "后端按角色校验模块权限，管理员、招生、教师、编辑等角色拥有不同访问范围。"),
    ]
    day2_code = r'''
# backend/app/api/auth.py
@router.post("/student-login")
def student_login(payload: StudentLoginRequest, db: Session = Depends(get_db)):
    student = db.query(StudentRecord).filter(StudentRecord.student_no == payload.student_no).first()
    if not student or payload.password != student.student_no:
        raise HTTPException(status_code=401, detail="学号或密码错误")
    token = create_access_token(
        f"student:{student.id}",
        {"role": "student", "student_no": student.student_no, "student_id": student.id},
    )
    return {"access_token": token, "student_name": student.name}

# frontend/src/views/StudentScheduleView.vue
const week = ref(1)
const schedule = computed(() =>
  allSchedules.value.filter(item => item.week_start <= week.value && item.week_end >= week.value)
)
'''

    day3_problems = [
        ("招生咨询如果只靠固定问答，无法回答分数、专业对比、就业方向等复杂问题。", "建设 AI 助手接口，结合本地专业库、分数线、官网采集内容和大模型生成回答。"),
        ("星火 API 因授权返回 11200，影响模型调用。", "新增智谱 GLM 接入，前端提供模型平台选择，默认使用已验证可用的智谱。"),
        ("AI 回答需要有依据，不能只给模型自由发挥。", "先生成本地检索结果，再把本地依据传入模型提示词，回答中保留 sources。"),
        ("校园生活百事通和招生咨询属于不同场景。", "通过 assistant_mode 区分 admissions 和 campus_life，分别走招生知识库和校园生活 RAG。"),
        ("用户希望回答能被读出来，提升展示效果。", "在前端加入数字人和 speechSynthesis 语音朗读，回答完成后自动播报。"),
    ]
    day3_code = r'''
# backend/app/services/ai_service.py
def _provider_config(payload: ChatRequest):
    if payload.provider == "zhipu":
        return settings.zhipu_api_key, settings.zhipu_base_url, settings.zhipu_model
    if payload.provider == "spark":
        return settings.spark_api_password, settings.spark_base_url, settings.spark_model
    if payload.provider == "deepseek":
        return settings.deepseek_api_key, settings.deepseek_base_url, settings.deepseek_model
    return settings.openai_api_key, settings.openai_base_url, settings.openai_model

def ask_ai(db: Session, payload: ChatRequest) -> ChatResponse:
    if payload.assistant_mode == "campus_life":
        local = campus_life_answer(payload)
    else:
        local = _local_answer(db, payload.question, _intent(payload.question))
    api_key, base_url, model = _provider_config(payload)
    client = OpenAI(api_key=api_key, base_url=base_url)
    result = client.chat.completions.create(model=model, messages=_spark_messages(payload, local))
'''

    day4_problems = [
        ("官网内容如果完全手动录入，维护成本高。", "实现官网采集模块，可配置采集源并抓取学校新闻、通知公告、招标采购等栏目。"),
        ("采集内容需要审核，不能直接发布到官网。", "采集文章可设置为待审核、草稿或已发布，并进入官网管理模块进行审核。"),
        ("采集失败时管理员需要知道原因。", "采集记录保存状态、原文链接、错误信息和生成文章 id，后台表格展示最近信息。"),
        ("官网页面需要从后台数据实时更新。", "前台文章、栏目、轮播图接口从数据库读取，后台修改后前台刷新即可看到。"),
        ("英文官网不能只是跳转空页面。", "增加英文官网页面和 English 入口，让右上角 English 能进入英文版展示。"),
    ]
    day4_code = r'''
# frontend/src/views/admin/CrawlerAdminView.vue
async function runOne(row) {
  runningId.value = row.id
  try {
    const { data } = await http.post(`/admin/crawl/run/${row.id}`, null, { params: { limit: 8 } })
    ElMessage.success(`${row.name} 采集完成：新增 ${data.created} 条，跳过 ${data.skipped} 条`)
    await loadAll()
  } finally {
    runningId.value = null
  }
}

# backend/app/api/admin.py
@router.post("/admin/crawl/run/{source_id}")
def run_crawl_source(source_id: int, limit: int = 8, db: Session = Depends(get_db)):
    source = db.get(CrawlSource, source_id)
    result = crawl_source(db, source, limit=limit)
    return result
'''

    day5_problems = [
        ("项目功能较多，最终需要保证整体可运行。", "运行后端语法检查、前端 build、接口测试和关键词乱码扫描。"),
        ("模型平台配置容易写死，后期切换困难。", "在 .env 中配置 ZHIPU、SPARK、DEEPSEEK、OPENAI，多平台统一由 provider 分流。"),
        ("数字人朗读可能在离开页面后继续播放。", "组件卸载时调用 stopSpeech()，并提供停止按钮。"),
        ("实训报告需要和实际项目一致。", "将报告内容从单一校园生活助手改为智慧官网与 AI 招生服务系统。"),
        ("部署时前端和后端需要明确启动方式。", "保留 run_backend.bat、run_frontend.bat 和 README 运行说明。"),
    ]
    day5_code = r'''
// frontend/src/views/ChatView.vue
function speakAnswer(text) {
  const content = cleanSpeechText(text)
  if (!voiceEnabled.value || !content || !('speechSynthesis' in window)) return
  stopSpeech()
  const utterance = new SpeechSynthesisUtterance(content)
  utterance.lang = 'zh-CN'
  utterance.rate = 1
  utterance.pitch = 1.05
  utterance.onstart = () => { isSpeaking.value = true }
  utterance.onend = () => { isSpeaking.value = false }
  window.speechSynthesis.speak(utterance)
}

onBeforeUnmount(() => stopSpeech())
'''

    tables = doc.tables
    fill_table(tables[0], "2026年6月17日", "敏行楼B417", "系统需求分析与前后端架构搭建", day1_problems, day1_code, "第 1 天主要完成项目整体规划和基础架构搭建。我把项目定位为智慧官网与 AI 招生服务系统，而不是单一问答程序。通过梳理官网门户、后台管理、招生服务、学生教务、官网采集和 AI 助手六条主线，确定了 Vue + FastAPI + MySQL 的前后端分离架构，并为后续功能扩展预留了路由、接口和数据模型。")
    fill_table(tables[1], "2026年6月18日", "敏行楼B417", "学生教务与后台管理模块完善", day2_problems, day2_code, "第 2 天重点完善后台和学校业务模块。我导入并整理学生数据，区分学生登录与管理员登录，加入学生课表、成绩、考勤和详情页联动，同时完善教务管理中的成绩总评计算和权限控制。这一天让我理解到校园系统不仅要有好看的页面，更要保证不同身份看到的数据范围正确。")
    fill_table(tables[2], "2026年6月19日", "敏行楼B417", "AI 招生助手、智谱模型与校园生活 RAG", day3_problems, day3_code, "第 3 天主要完成 AI 能力建设。我将招生咨询、专业推荐、校园生活百事通统一到 /chat 页面中，通过 assistant_mode 区分业务场景，通过 provider 支持智谱、星火、DeepSeek 和 OpenAI。由于星火授权受限，我接入并验证了智谱 GLM，使系统能够稳定调用大模型。")
    fill_table(tables[3], "2026年6月20日", "敏行楼B417", "官网实时采集与内容管理", day4_problems, day4_code, "第 4 天完成官网采集和内容管理优化。系统支持配置采集源、运行单个或全部采集任务、查看采集记录，并将采集内容写入官网文章库进行审核。这样官网内容不再只是静态页面，而是可以通过后台持续维护和更新，提升了系统的真实性和完整度。")
    fill_table(tables[4], "2026年6月21日", "敏行楼B417", "数字人播报、测试与交付整理", day5_problems, day5_code, "第 5 天进行项目收尾和展示增强。我为招生 AI 助手加入数字人播报功能，回答生成后自动朗读，并提供重播、停止和静音控制。同时运行前端打包、后端检查和模型接口测试，确认智谱 GLM 可用。最后整理 README、测试报告和实训报告，使项目具备完整演示和提交条件。")

    if output.exists():
        output.unlink()
    doc.save(str(output))
    print(output)


if __name__ == "__main__":
    main()
