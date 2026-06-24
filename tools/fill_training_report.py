from pathlib import Path
import os

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def set_font(run, size=11, bold=False, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_para_text(paragraph, text, size=11, bold=False, align=None):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    if align is not None:
        paragraph.alignment = align


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, size=10.5, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def fill_day_table(table, date_text, place, task, problems, code, summary):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_text(table.cell(0, 1), date_text)
    cell_text(table.cell(0, 3), place)
    cell_text(table.cell(1, 1), task)
    for col in range(4):
        shade_cell(table.cell(0, col), "EAF2F8")
        shade_cell(table.cell(1, col), "F6F8FA")

    for index, row_index in enumerate([4, 5, 6, 7, 8]):
        if index < len(problems):
            cell_text(table.cell(row_index, 0), problems[index][0])
            cell_text(table.cell(row_index, 2), problems[index][1])
        else:
            cell_text(table.cell(row_index, 0), "")
            cell_text(table.cell(row_index, 2), "")

    cell_text(table.cell(10, 0), code, size=9)
    cell_text(table.cell(len(table.rows) - 1, 0), summary)


def main():
    desktop = Path(os.environ.get("USERPROFILE", r"C:\Users\30587")) / "Desktop"
    source = next(path for path in desktop.glob("202431212*.docx") if " - " in path.name)
    output = desktop / "202431212  王振宇 实训报告 - 完成版.docx"
    doc = Document(str(source))

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    replacements = {
        "项目名称：__________________": "项目名称：校园生活百事通助手",
        "学生姓名：______ 学号：______": "学生姓名：王振宇    学号：202431212",
        "班级：_24人工智能_ 指导教师：_魏化永、单列_": "班级：24人工智能    指导教师：魏化永、单列",
        "实训日期：2026年__月__日 - __月__日": "实训日期：2026年6月17日 - 2026年6月21日",
    }

    paragraphs = doc.paragraphs
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_para_text(paragraph, replacements[text], size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif text == "1. 项目概述（200字）":
            set_para_text(paragraph, text, size=14, bold=True)
        elif text == "项目名称、解决的问题、主要功能":
            set_para_text(
                paragraph,
                "本项目名称为“校园生活百事通助手”，面向在校学生日常办事咨询场景，解决学生在请假、奖学金、宿舍报修、一卡通补办、选课退课等问题上信息分散、查询效率低、容易遗漏流程要求的问题。系统基于现有 FastAPI + Vue 前后端分离项目实现，在原有 AI 助手中增加“校园生活”模式。后端整理 20 条校园生活规则问答，保存为 campus_data.csv，并构建 campus_life_kb 向量知识库；前端提供聊天界面、常用问题入口、模型平台切换和参考来源展示。系统支持 RAG 检索增强回答、教学周查询、绩点计算、多轮对话记忆和低置信度兜底提示，能够较完整地完成案例 A 的开发目标。",
            )
        elif text == "2. 技术架构图":
            set_para_text(paragraph, text, size=14, bold=True)
        elif text == "（手绘或流程图）":
            set_para_text(
                paragraph,
                "技术架构图：用户 -> Vue 聊天界面(/chat) -> FastAPI /api/chat 接口 -> 校园生活智能体 -> 意图识别 -> RAG 检索/工具调用 -> Chroma 向量库(campus_life_kb)、CSV 知识库(campus_data.csv)、绩点计算工具、教学周工具 -> 生成回答并返回参考来源。",
            )
        elif text == "画出：用户→界面→智能体→RAG/工具→向量库/数据库":
            set_para_text(
                paragraph,
                "流程说明：用户在页面选择“校园生活”模式并输入问题；前端将 assistant_mode、问题和最近历史发送到后端；后端先判断是否为教学周或绩点计算工具请求，若不是则进入 RAG 流程，从校园生活向量库和 CSV 规则库中检索相关片段，再根据提示词要求生成简洁、友好的回答；若无法检索到依据，则返回“我不确定，建议咨询辅导员”。",
            )
        elif text == "4. 核心代码展示":
            set_para_text(paragraph, text, size=14, bold=True)
        elif text == "粘贴最重要的3个函数（带注释）":
            set_para_text(
                paragraph,
                "核心函数1：retrieve_campus_knowledge(question, limit=3)\n作用：对用户问题进行同义词扩展，并优先使用 campus_life_kb 向量库检索相关规则；若向量库不可用，则回退到关键词和相似度重排，保证系统稳定可用。\n\n核心函数2：campus_life_answer(payload)\n作用：实现智能体入口。先判断“现在第几周”“绩点计算”等工具意图；普通校园问题进入 RAG 检索，按规则库生成答案；未知问题不编造，提示咨询辅导员。\n\n核心函数3：build_campus_vector_db()\n作用：读取 campus_data.csv 中的 20 条问答，转换为“问题+答案”的文本片段，写入 Chroma 的 campus_life_kb 集合，供 RAG 检索使用。",
            )
        elif text == "5. 运行截图":
            set_para_text(paragraph, text, size=14, bold=True)
        elif text == "至少3张不同功能的截图":
            set_para_text(
                paragraph,
                "截图说明：\n1. /chat 页面切换到“校园生活”模式，展示常用问题入口。\n2. 输入“怎么请病假？”后，系统返回包含辅导员、请假申请、校医院证明等内容的规则库回答，并显示参考来源。\n3. 输入“绩点计算 85,90,78”后，系统调用绩点工具，返回平均绩点 3.00。\n4. 输入“现在第几周？”后，系统调用教学周工具，返回当前教学周估算结果。\n实际提交 PDF 时可将以上运行页面截图粘贴到本节下方。",
            )
        elif text == "6. 问题与反思（300字）":
            set_para_text(paragraph, text, size=14, bold=True)
        elif text == "遇到的最大困难是什么？":
            set_para_text(
                paragraph,
                "本次实训中遇到的最大困难是如何把指导书中的独立 Streamlit 示例融合到已有的完整网站系统中。原案例按照 data、src、vector_db 的单体项目结构展开，而我的项目已经使用 FastAPI 后端、Vue 前端、MySQL 数据库和现有 AI 聊天页面。如果完全照搬示例，会形成两个互不关联的系统。因此我选择保留指导书的 RAG 思路和智能体能力，把校园生活助手作为现有 /chat 页面中的一种模式。另一个困难是中文编码和检索准确率。Windows 环境下如果写入方式不正确，中文容易变成问号；检索方面，简单向量哈希对“我发烧了怎么办”这类口语表达不够敏感。最后通过统一 UTF-8 写入、增加同义词扩展和重排策略解决。还有一些功能可以继续完善，例如接入学校真实学生手册、办事大厅接口和更强的中文嵌入模型。如果再做一次，我会先设计标准数据表和后台维护页面，让老师或管理员可以直接在网页上增删校园规则，而不是只通过 CSV 文件维护。",
            )
        elif text in {"哪些功能想做但没做完？", "如果再有一次机会，会改进什么？"}:
            set_para_text(paragraph, "")
        elif text == "7. 实训总结（200字）":
            set_para_text(paragraph, text, size=14, bold=True)
        elif text == "学到了什么？对AI开发的理解变化":
            set_para_text(
                paragraph,
                "通过本次实训，我完整体验了一个 AI 应用从需求分析、数据整理、知识库构建、检索增强、工具调用到 Web 集成的过程。过去我对大模型应用的理解主要停留在“向模型提问并得到回答”，这次实践让我认识到，真正可用的 AI 系统不能只依赖模型本身，还需要可靠的数据来源、清晰的提示词、可控的工具函数和完善的异常兜底。校园生活百事通助手虽然规模不大，但包含了 RAG 的核心流程：先检索规则，再增强上下文，最后生成回答。同时，绩点计算和教学周查询也让我理解了智能体工具调用的意义。通过把案例 A 融合进已有的前后端系统，我进一步熟悉了 FastAPI 接口、Vue 页面交互、Chroma 向量库和配置管理。总体来看，本次实训提升了我把 AI 技术落地到实际校园服务场景中的能力。",
            )

    fill_day_table(
        doc.tables[0],
        "2026年6月17日",
        "实训室/宿舍",
        "数据收集与结构化整理",
        [
            ("校园生活问题来源比较分散，请假、奖学金、报修等规则不在同一位置。", "按照指导书要求归纳为 5 类必选主题，并扩展到 20 条有效问答。"),
            ("问答数据字段不统一，后续不便检索。", "统一整理为 id、category、question、answer、source 五个字段。"),
            ("部分规则不是学校实时官方数据。", "在答案中加入“以学校最新通知和辅导员要求为准”的提示。"),
        ],
        "data/campus_data.csv：保存 20 条校园生活问答数据。\n字段：id, category, question, answer, source。",
        "第 1 天主要完成了项目选题和数据准备。我围绕校园生活中学生最常遇到的问题进行整理，重点覆盖请假流程、奖学金评定、宿舍报修、一卡通补办和选课退课五类主题，并继续扩展到考试、图书馆、校园网、在读证明、就业信息等场景。通过把这些内容整理成 CSV 文件，我理解了 RAG 系统中“数据质量决定回答质量”的重要性。",
    )
    fill_day_table(
        doc.tables[1],
        "2026年6月18日",
        "实训室/宿舍",
        "向量库构建与检索测试",
        [
            ("原指导书使用 BAAI/bge-small-zh，首次下载可能受网络影响。", "结合现有项目使用 Chroma + HashEmbeddingFunction，保证本地可运行。"),
            ("“我发烧了怎么办”一开始没有稳定命中请假流程。", "增加“发烧=病假/请假/校医院/诊断证明”的同义词扩展。"),
            ("向量检索结果需要和规则库答案对应。", "为每条向量数据保存 category、question、source 元数据。"),
        ],
        "build_campus_vector_db()：读取 CSV，写入 campus_life_kb。\nretrieve_campus_knowledge()：检索并重排相关校园规则。",
        "第 2 天重点完成了向量库建立和检索测试。我将 campus_data.csv 中的 20 条问答转换为“问题+答案”的文本片段，并写入 Chroma 的 campus_life_kb 集合。测试中发现口语化表达会影响命中效果，因此补充了同义词扩展和重排逻辑。最终“我发烧了怎么办”能够返回请假相关内容，“奖学金要多少绩点”能够返回 3.0 条件，达到了检索检查点要求。",
    )
    fill_day_table(
        doc.tables[2],
        "2026年6月19日",
        "实训室/宿舍",
        "RAG 问答系统实现",
        [
            ("如果只调用大模型，可能会编造不存在的学校规定。", "设置系统提示词，要求只根据校园规则回答，不确定则建议咨询辅导员。"),
            ("本地没有配置 API 时也要能运行。", "无 API 时返回本地规则库结果，有 API 时再调用 Spark 或 DeepSeek 生成。"),
            ("回答需要展示依据。", "返回 sources 字段，前端显示参考来源和规则摘要。"),
        ],
        "campus_life_answer(payload)：根据问题检索规则库，组织回答和来源。\nRAG_PROMPT：限制回答范围，避免编造。",
        "第 3 天完成了 RAG 问答主流程。我把检索到的校园规则作为回答依据，要求助手只根据规则回答。对于未知问题，系统会提示“我不确定，建议咨询辅导员”，避免生成虚假制度。通过测试“怎么请病假”“一卡通丢了怎么办”“选错了课能退吗”等问题，系统能够返回较准确的流程说明，并附带来源信息。",
    )
    fill_day_table(
        doc.tables[3],
        "2026年6月20日",
        "实训室/宿舍",
        "智能体工具与多轮对话",
        [
            ("绩点问题既可能是奖学金条件，也可能是计算工具请求。", "只有“绩点计算”或带具体分数时才触发 GPA 工具。"),
            ("教学周计算需要有明确学期起点。", "按 2026 年 3 月 2 日作为本学期起点进行估算，并提示以校历为准。"),
            ("聊天页需要保留上下文。", "前端提交最近 10 条历史消息，后端接收 history 字段。"),
        ],
        "get_current_week()：返回当前教学周。\ncalculate_gpa(scores)：按 90/80/70/60 分段估算绩点。",
        "第 4 天主要完成智能体工具调用。系统能够识别“现在第几周”并调用教学周工具，也能识别“绩点计算 85,90,78”并返回平均绩点 3.00。开发中我发现“奖学金要多少绩点”不能被误判成绩点计算，因此对意图识别进行了细化。通过这一步，我理解了智能体并不只是聊天，还可以根据用户意图选择合适工具完成确定性任务。",
    )
    fill_day_table(
        doc.tables[4],
        "2026年6月21日",
        "实训室/宿舍",
        "系统集成、测试与文档整理",
        [
            ("指导书使用 Streamlit，但当前项目是 Vue + FastAPI。", "不再新增 Streamlit，将案例 A 集成到现有 /chat 页面。"),
            ("前端需要兼容多个模型平台。", "增加 Spark-X2、DeepSeek、OpenAI 兼容平台选择。"),
            ("需要形成可提交材料。", "补充测试报告、README 说明、演示脚本和构建脚本。"),
        ],
        "ChatView.vue：增加“招生咨询/校园生活”模式切换。\nai_service.py：根据 assistant_mode 选择校园生活或招生咨询逻辑。",
        "第 5 天完成了系统集成和最终测试。我在 Vue 聊天页中加入“校园生活”模式，并保留原有招生咨询功能。系统支持 Spark-X2、DeepSeek 和 OpenAI 兼容平台选择，同时在没有 API 密钥时也能基于本地规则库回答。最后运行了后端语法检查、前端打包、向量库构建和检索测试，并整理 README 与测试报告。本项目已经能够作为校园生活百事通助手进行演示。",
    )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_font(run, size=9.5 if len(cell.text) > 80 else 10.5)

    if output.exists():
        output.unlink()
    doc.save(str(output))
    print(output)


if __name__ == "__main__":
    main()
