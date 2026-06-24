from pathlib import Path
import os

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font(run, size=9.5, name="宋体", bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def write_cell(cell, text, size=10, bold=False, font_name="宋体"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=size, name=font_name, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def write_code_cell(cell, text):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    for index, line in enumerate(text.strip("\n").splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=8.2, name="Consolas")


def fill_problems(table, rows):
    for index, row_index in enumerate([4, 5, 6, 7, 8]):
        if index < len(rows):
            write_cell(table.cell(row_index, 0), rows[index][0], size=9.5)
            write_cell(table.cell(row_index, 2), rows[index][1], size=9.5)
        else:
            write_cell(table.cell(row_index, 0), "", size=9.5)
            write_cell(table.cell(row_index, 2), "", size=9.5)


def main():
    desktop = Path(os.environ.get("USERPROFILE", r"C:\Users\30587")) / "Desktop"
    source = next(path for path in desktop.glob("202431212*.docx") if "完成版" in path.name and "每日记录完善版" not in path.name)
    output = desktop / "202431212  王振宇 实训报告 - 每日记录完善版.docx"
    doc = Document(str(source))

    day1_problems = [
        (
            "校园生活规则来源分散，官网、学生手册、班级群通知的表述方式不一致，直接拿来做知识库会导致回答混乱。",
            "先按指导书要求确定 5 个核心主题：请假、奖学金、宿舍报修、一卡通、选课退课，再统一整理成“问题-答案-来源”的结构化格式。",
        ),
        (
            "部分校园规则带有时效性，例如奖学金评定、选课退课时间每学期可能不同，不能直接写成绝对结论。",
            "在答案中保留流程性说明，并补充“以学校最新通知、辅导员或教务部门要求为准”的提示，降低误导风险。",
        ),
        (
            "原始问题表达比较口语化，例如“发烧了怎么办”“灯坏了找谁”，和正式制度标题差异较大。",
            "整理 CSV 时同时保留学生常见问法，让后续检索既能匹配正式类别，也能匹配真实提问习惯。",
        ),
        (
            "CSV 文件如果字段不统一，后续构建向量库时很难保存元数据，也不方便展示参考来源。",
            "统一字段为 id、category、question、answer、source，并要求每条记录都有明确分类和来源说明。",
        ),
        (
            "数据量既要满足指导书“至少 20 条”的要求，又不能为了凑数加入无关内容。",
            "围绕学生高频生活场景扩展到 20 条，包括助学、考试、图书馆、校园网、证明、就业、安全等内容。",
        ),
    ]
    day1_code = r'''
# backend/data/campus_data.csv 的核心字段
id,category,question,answer,source
1,请假,怎么请病假？,病假应先联系辅导员说明情况，填写请假申请，准备校医院或正规医院诊断证明。,学生手册示例：请假流程
5,奖学金,奖学金要多少绩点？,示例规则中，奖学金参考条件为学年平均绩点不低于 3.0、无挂科、体测合格。,学生手册示例：奖学金绩点
7,宿舍,宿舍灯坏了找谁？,宿舍设施损坏可先向宿管登记，或通过学校后勤报修渠道提交房间号和联系方式。,后勤报修流程示例

# backend/app/services/campus_life_service.py
def load_campus_knowledge() -> list[dict]:
    rows = []
    with CAMPUS_DATA_PATH.open("r", encoding="utf-8-sig", newline="") as file:
        for row in csv.DictReader(file):
            if not row.get("question") or not row.get("answer"):
                continue
            rows.append({
                "id": int(row.get("id") or len(rows) + 1),
                "category": row.get("category", "校园生活"),
                "question": row["question"],
                "answer": row["answer"],
                "source": row.get("source", "campus_data.csv"),
            })
    return rows or FALLBACK_CAMPUS_KNOWLEDGE
'''

    day2_problems = [
        (
            "指导书推荐使用 bge-small-zh 嵌入模型，但本地首次运行可能需要下载模型，网络不稳定会影响实训演示。",
            "复用项目已有的 Chroma 向量服务，并使用 HashEmbeddingFunction 做本地轻量嵌入，保证无网络时也能构建和检索。",
        ),
        (
            "Chroma 集合如果和招生知识库混在一起，会导致校园生活问题检索到招生内容。",
            "新增独立集合 campus_life_kb，把校园生活知识与招生知识分开存储和查询。",
        ),
        (
            "向量库重复构建时，旧数据可能残留，导致同一条规则出现多次。",
            "构建时使用 reset=True，先删除相同 id，再写入新的 20 条规则，保证结果干净。",
        ),
        (
            "“我发烧了怎么办？”第一次检索时没有稳定命中“病假”规则，说明口语表达和制度关键词之间有距离。",
            "增加 query expansion，把“发烧、生病、不舒服”扩展为“病假、请假、校医院、诊断证明”，再检索和重排。",
        ),
        (
            "只看向量距离不一定准确，短文本和中文口语容易出现误召回。",
            "把向量检索结果与关键词相似度结果合并，再用 _score() 统一重排，提升命中稳定性。",
        ),
    ]
    day2_code = r'''
# backend/app/services/vector_service.py
def chroma_collection(name: str = "admission_kb"):
    import chromadb
    VECTOR_DIR.mkdir(parents=True, exist_ok=True)
    client = chromadb.PersistentClient(path=str(VECTOR_DIR))
    return client.get_or_create_collection(name, embedding_function=HashEmbeddingFunction())

def add_chunks_to_chroma(chunks, collection_name="admission_kb", reset=False):
    collection = chroma_collection(collection_name)
    ids = [str(item["id"]) for item in chunks]
    if reset:
        try:
            collection.delete(ids=ids)
        except Exception:
            pass
    collection.add(
        ids=ids,
        documents=[item["content"] for item in chunks],
        metadatas=[{
            "source": item.get("source", ""),
            "category": item.get("category", ""),
            "question": item.get("question", ""),
        } for item in chunks],
    )

# backend/app/services/campus_life_service.py
def build_campus_vector_db() -> int:
    chunks = [{
        "id": f"campus-{item['id']}",
        "content": f"问题：{item['question']}\n答案：{item['answer']}",
        "source": item["source"],
        "category": item["category"],
        "question": item["question"],
    } for item in load_campus_knowledge()]
    add_chunks_to_chroma(chunks, collection_name="campus_life_kb", reset=True)
    return len(chunks)
'''

    day3_problems = [
        (
            "如果直接把用户问题交给大模型，模型可能会补充不存在的学校制度，出现“编造规则”的风险。",
            "设计校园生活系统提示词，明确要求只能依据规则库回答，不确定时必须建议咨询辅导员。",
        ),
        (
            "RAG 回答不能只返回结论，还要能说明依据，否则用户难以判断可信度。",
            "ChatResponse 中保留 sources 列表，前端展示参考来源和预览片段。",
        ),
        (
            "本地没有配置 Spark 或 DeepSeek API 时，AI 功能不能完全失效。",
            "先生成本地规则库答案；如果没有 API Key，直接返回本地答案；如果有 API，再用模型润色。",
        ),
        (
            "奖学金、选课等问题既有流程性又有时效性，回答需要避免过度确定。",
            "在回答末尾统一加入“具体办理以学校最新通知和辅导员要求为准”的提醒。",
        ),
        (
            "同一个问题可能检索到多条相关规则，例如病假和事假都属于请假。",
            "取前 2 条高相关规则组织答案，并把前 3 条作为 sources 返回，兼顾简洁和可追溯。",
        ),
    ]
    day3_code = r'''
# backend/app/services/ai_service.py
def _spark_messages(payload: ChatRequest, local: ChatResponse) -> list[dict[str, str]]:
    if payload.assistant_mode == "campus_life":
        system_prompt = (
            "你是安徽交通职业技术学院校园生活百事通助手。"
            "只能根据本地校园生活规则库和工具结果回答学生问题。"
            "如果规则里没有相关信息，要说“我不确定，建议咨询辅导员”。"
            "回答要简洁、友好，不要编造学校政策。"
        )
    ...

# backend/app/services/campus_life_service.py
def campus_life_answer(payload: ChatRequest) -> ChatResponse:
    question = payload.question.strip()
    hits = retrieve_campus_knowledge(question)
    if not hits:
        return ChatResponse(
            intent="校园生活问答",
            answer="我不确定，建议咨询辅导员或相关部门。",
            confidence=0.42,
            sources=[],
        )
    lines = ["根据校园生活规则库，可以这样处理：", ""]
    for item in hits[:2]:
        lines.append(f"- {item['answer']}")
    lines.append("\n提醒：以上为校园生活百事通助手的规则库回答，具体办理以学校最新通知和辅导员要求为准。")
    sources = [{"source": item["source"], "preview": f"{item['category']}：{item['question']}"} for item in hits]
    return ChatResponse(intent="校园生活问答", answer="\n".join(lines), confidence=0.82, sources=sources)
'''

    day4_problems = [
        (
            "“奖学金要多少绩点”和“绩点计算 85,90,78”都包含“绩点”，一开始容易误触发绩点计算工具。",
            "将触发条件改为“绩点计算”或包含具体分数时才调用 GPA 工具，奖学金问题继续走规则库。",
        ),
        (
            "教学周计算必须有明确学期起始日期，否则每次回答会不一致。",
            "设置 2026 年 3 月 2 日为本学期第 1 周起点，并在答案中说明是估算值。",
        ),
        (
            "用户没有输入分数却要求计算绩点时，工具无法计算。",
            "返回提示“请告诉我各科分数，例如：绩点计算 85,90,78”，引导用户补全输入。",
        ),
        (
            "多轮对话如果不传历史，模型无法知道前面聊过什么。",
            "前端 recentHistory() 保留最近 10 条消息，随请求一起发送到后端。",
        ),
        (
            "工具回答和 RAG 回答的来源展示格式不同。",
            "工具调用也返回 sources，例如“校园生活工具：按 90/80/70/60 分段规则估算平均绩点”。",
        ),
    ]
    day4_code = r'''
# backend/app/services/campus_life_service.py
def get_current_week(today: date | None = None) -> str:
    today = today or date.today()
    term_start = date(2026, 3, 2)
    week_num = (today - term_start).days // 7 + 1
    if week_num < 1:
        return "当前还未到本学期第 1 周，具体校历请以教务通知为准。"
    return f"按 2026 年 3 月 2 日开学估算，现在是第 {week_num} 周。"

def calculate_gpa(scores: list[int]) -> str:
    if not scores:
        return "请告诉我各科分数，例如：绩点计算 85,90,78。"
    points = []
    for score in scores:
        if score >= 90: points.append(4.0)
        elif score >= 80: points.append(3.0)
        elif score >= 70: points.append(2.0)
        elif score >= 60: points.append(1.0)
        else: points.append(0.0)
    return f"按示例规则估算，您的平均绩点是：{sum(points) / len(points):.2f}。"

scores = [int(num) for num in re.findall(r"\d+", question) if 0 <= int(num) <= 100]
wants_gpa_tool = "GPA" in question.upper() or "绩点计算" in question or ("绩点" in question and bool(scores))
if wants_gpa_tool:
    return ChatResponse(intent="校园工具：绩点计算", answer=calculate_gpa(scores), confidence=0.92, sources=[...])
'''

    day5_problems = [
        (
            "指导书使用 Streamlit，但当前项目已经是完整的 Vue + FastAPI 系统，如果另做页面会割裂。",
            "将案例 A 集成到现有 /chat 页面，通过 assistant_mode 在招生咨询和校园生活之间切换。",
        ),
        (
            "不同模型平台的 API 配置不一样，系统原来主要面向 Spark-X2。",
            "增加 provider 选择，支持 Spark-X2、DeepSeek、OpenAI 兼容平台，并在后端统一 _provider_config()。",
        ),
        (
            "前端用户需要知道当前使用的是哪种助手能力。",
            "在页面顶部和状态面板显示“案例 A：校园百事通”“RAG 规则库 + 工具调用”。",
        ),
        (
            "实训提交不仅要代码，还需要可运行脚本和测试记录。",
            "补充 build_campus_vector_db.py、test_campus_retrieve.py、campus_rag_demo.py、campus_agent_demo.py 和测试报告。",
        ),
        (
            "项目多次编辑后容易出现中文编码损坏。",
            "最终扫描源码中的明显 ???? 乱码，并运行前端打包、后端语法检查确认可用。",
        ),
    ]
    day5_code = r'''
# frontend/src/views/ChatView.vue
const assistantMode = ref('admissions')
const provider = ref('spark')
const modeOptions = [
  { label: '招生咨询', value: 'admissions' },
  { label: '校园生活', value: 'campus_life' }
]
function payload(q, history) {
  return {
    question: q,
    assistant_mode: assistantMode.value,
    provider: provider.value,
    thinking: thinking.value,
    web_search: webSearch.value,
    search_mode: searchMode.value,
    history
  }
}

# backend/app/services/ai_service.py
def _provider_config(payload: ChatRequest):
    if payload.provider == "deepseek":
        return settings.deepseek_api_key, settings.deepseek_base_url, settings.deepseek_model
    if payload.provider == "spark":
        return settings.spark_api_password, settings.spark_base_url, settings.spark_model
    return settings.openai_api_key, settings.openai_base_url, settings.openai_model

def ask_ai(db: Session, payload: ChatRequest) -> ChatResponse:
    if payload.assistant_mode == "campus_life":
        local = campus_life_answer(payload)
        ...
'''

    tables = doc.tables
    fill_problems(tables[0], day1_problems)
    write_code_cell(tables[0].cell(10, 0), day1_code)
    fill_problems(tables[1], day2_problems)
    write_code_cell(tables[1].cell(10, 0), day2_code)
    fill_problems(tables[2], day3_problems)
    write_code_cell(tables[2].cell(10, 0), day3_code)
    fill_problems(tables[3], day4_problems)
    write_code_cell(tables[3].cell(10, 0), day4_code)
    fill_problems(tables[4], day5_problems)
    write_code_cell(tables[4].cell(10, 0), day5_code)

    if output.exists():
        output.unlink()
    doc.save(str(output))
    print(output)


if __name__ == "__main__":
    main()
